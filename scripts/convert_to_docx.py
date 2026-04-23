"""Convert Antrag markdown files to DOCX in KDV Beschlussbuch format."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_margins(doc):
    """Set A4 page size and standard margins."""
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def add_header(doc, antrag_title):
    """Add header matching Beschlussbuch style."""
    header = doc.sections[0].header
    p = header.paragraphs[0]
    p.text = ""
    run = p.add_run("Anträge · KDV am 18.04.2026")
    run.font.size = Pt(9)
    run.font.name = "Arial"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Add bottom border to header paragraph
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_footer(doc):
    """Add centered page number footer."""
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add top border
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "000000")
    pBdr.append(top)
    pPr.append(pBdr)
    # Page number field
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run1 = p.add_run()
    run1._element.append(fld_char_begin)
    run1.font.size = Pt(9)
    run1.font.name = "Arial"
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run2 = p.add_run()
    run2._element.append(instr)
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run3 = p.add_run()
    run3._element.append(fld_char_end)


def add_voting_table(doc):
    """Add empty voting table at the bottom."""
    doc.add_paragraph()  # spacer

    table = doc.add_table(rows=4, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    cells = table.rows[0].cells
    cells[0].text = "Votum Antragskommission"
    cells[0].merge(cells[1])
    cells[2].text = "Abstimmung KDV"
    cells[2].merge(cells[3])

    # Make headers italic
    for cell in [cells[0], cells[2]]:
        for p in cell.paragraphs:
            for run in p.runs:
                run.italic = True
                run.font.size = Pt(9)
                run.font.name = "Arial"

    # Option rows
    options = [
        ("Änderungen", "Überweisung", "Überweisung", ""),
        ("Zustimmung", "Ablehnung", "Änderung", ""),
        ("Konsens", "mehrheitlich", "Zustimmung", "Ablehnung"),
    ]
    for i, (a, b, c, d) in enumerate(options):
        row = table.rows[i + 1]
        row.cells[0].text = a
        row.cells[1].text = b
        row.cells[2].text = c
        row.cells[3].text = d

    # Format all cells
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Arial"


def suppress_line_numbers(paragraph):
    """Suppress line numbering for a specific paragraph."""
    pPr = paragraph._element.get_or_add_pPr()
    suppress = OxmlElement("w:suppressLineNumbers")
    pPr.append(suppress)


def add_run_with_formatting(paragraph, text):
    """Add text to paragraph, handling **bold** and \\*gendering."""
    # Split on bold markers
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Replace \* with * for gendering
            clean = part.replace("\\*", "*")
            paragraph.add_run(clean)


def md_to_docx(md_path: Path, out_path: Path):
    """Convert a single markdown Antrag to DOCX in KDV format."""
    text = md_path.read_text(encoding="utf-8")

    doc = Document()

    # Page setup
    set_margins(doc)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(10)
    font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    lines = text.split("\n")

    # Extract title (first # line)
    title = ""
    for line in lines:
        if line.strip().startswith("# "):
            title = line.strip()[2:]
            break

    # Add header/footer
    add_header(doc, title)
    add_footer(doc)

    # Title — bold, no heading, no line numbers
    p = doc.add_paragraph()

    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Arial"
    p.paragraph_format.space_after = Pt(6)

    # Add KDV intro line — italic, no line numbers
    p = doc.add_paragraph()

    run = p.add_run("Die Kreisdelegiertenversammlung wolle beschließen:")
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = "Arial"
    p.paragraph_format.space_after = Pt(2)

    # Parse content
    skip_title = True
    in_begruendung = False
    found_adressat = False

    for line in lines:
        stripped = line.strip()

        # Skip empty lines — but add spacing
        if not stripped:
            continue

        # Skip the markdown title
        if stripped.startswith("# ") and skip_title:
            skip_title = False
            continue

        # Skip ## Begründung heading — replace with bold inline
        if stripped == "## Begründung":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            run = p.add_run("Begründung:")
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = "Arial"
            in_begruendung = True
            continue

        # Skip other ## headings
        if stripped.startswith("## "):
            continue

        # Adressat line (contains "mögen beschließen" or "möge beschließen")
        if not found_adressat and ("mögen beschließen" in stripped or "möge beschließen" in stripped):
            p = doc.add_paragraph()
        
            # Clean markdown bold
            clean = stripped.replace("**", "")
            run = p.add_run(clean)
            run.italic = True
            run.font.size = Pt(10)
            run.font.name = "Arial"
            p.paragraph_format.space_after = Pt(6)
            found_adressat = True
            continue

        # Numbered list items (Forderungen)
        if re.match(r"^\d+\.\s", stripped):
            content = stripped  # Keep the number
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            add_run_with_formatting(p, content)
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = "Arial"
            continue

        # Bullet list items
        if stripped.startswith("- "):
            content = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            bullet_run = p.add_run("•  ")
            bullet_run.font.size = Pt(10)
            bullet_run.font.name = "Arial"
            add_run_with_formatting(p, content)
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = "Arial"
            continue

        # Table rows — skip
        if stripped.startswith("|"):
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_run_with_formatting(p, stripped)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = "Arial"

    doc.save(str(out_path))
    print(f"  {out_path.name}")


def main():
    base = Path(__file__).parent
    md_files = sorted(f for f in base.glob("*.md") if f.name != "convert_to_docx.py")

    # Skip the internal file
    md_files = [f for f in md_files if "INTERN" not in f.name]

    print(f"Converting {len(md_files)} Anträge to DOCX (KDV format):\n")
    for md in md_files:
        out = base / md.name.replace(".md", ".docx")
        md_to_docx(md, out)

    print(f"\nDone. {len(md_files)} files written.")


if __name__ == "__main__":
    main()
